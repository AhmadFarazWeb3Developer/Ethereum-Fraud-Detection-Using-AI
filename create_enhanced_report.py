import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir(r'c:\Users\MMO\desktop\Ethereum-Fraud-Detection-Using-AI\Assignment_02')

def add_table_of_contents(doc):
    """Add a table of contents to the document"""
    pass  # Word will generate this automatically

def shade_cell(cell, color):
    """Shade a cell with a specific color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Assignment 02: Exploratory Data Analysis (EDA)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Course Information
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'
info_table.cell(0, 0).text = 'Course'
info_table.cell(0, 1).text = 'Data Science'
info_table.cell(1, 0).text = 'Class'
info_table.cell(1, 1).text = 'BSCS-F22'
info_table.cell(2, 0).text = 'Instructor'
info_table.cell(2, 1).text = 'Mr. Ghulam Ali'
info_table.cell(3, 0).text = 'Student Name'
info_table.cell(3, 1).text = 'Ahmad Faraz'
info_table.cell(4, 0).text = 'Registration No'
info_table.cell(4, 1).text = '215154'

# Header cells shading
for i in range(2):
    shade_cell(info_table.cell(0, i), 'D3D3D3')
    shade_cell(info_table.cell(1, i), 'D3D3D3')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Introduction",
    "2. Dataset Loading and Overview",
    "3. Data Shape and Structure",
    "4. Descriptive Statistics",
    "5. Target Variable Analysis (FLAG)",
    "6. Class Distribution Visualization",
    "7. Feature Distribution Analysis",
    "8. Outlier Detection using Boxplots",
    "9. Correlation Analysis",
    "10. Relationship Analysis",
    "11. Key Findings and Insights",
    "12. Conclusion"
]

for item in toc_items:
    toc_para = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
intro_text = """This notebook builds upon Assignment 01, where Ethereum blockchain fraud data was collected and preprocessed. The primary goal of this assignment is to perform comprehensive Exploratory Data Analysis (EDA), extract meaningful insights from the data, and establish a methodological foundation for fraud detection.

The dataset contains transactions from the Ethereum blockchain with various features that help identify fraudulent activities. Through this EDA, we aim to:
• Understand the structure and characteristics of the data
• Identify patterns and anomalies in blockchain transactions
• Explore relationships between features
• Detect outliers and unusual transactions
• Build a foundation for predictive modeling"""
doc.add_paragraph(intro_text)

# 2. Dataset Loading and Overview
doc.add_heading('2. Dataset Loading and Overview', level=1)
doc.add_paragraph("""The cleaned Ethereum fraud detection dataset is loaded using pandas. This dataset contains preprocessed transactions with no missing values after the data cleaning process performed in Assignment 01.""")

# Add code snippet
code_para = doc.add_paragraph()
code_para.style = 'Normal'
code_run = code_para.add_run("""import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

df = pd.read_csv('Cleaned_Ethereum_Fraud_Detection.csv')
df.head()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# 3. Data Shape and Structure
doc.add_heading('3. Data Shape and Structure', level=1)
doc.add_paragraph("""The dataset dimensions and column information provide insights into the data structure. The shape tells us how many transactions and features we have, while the info provides details about data types and memory usage.""")

# Add code snippet
code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.shape
df.info()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

doc.add_paragraph("""Expected Output:
• Total rows: 9,841 transactions
• Total columns: 48 features
• Data types: Mix of integer, float, and potentially categorical features
• Memory usage: Approximately 3.7+ MB""")

# 4. Descriptive Statistics
doc.add_heading('4. Descriptive Statistics', level=1)
doc.add_paragraph("""Descriptive statistics provide a summary of the numerical features in the dataset, including measures of central tendency and spread.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.describe()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

doc.add_paragraph("""This output includes:
• count: Number of non-null observations
• mean: Average value of each feature
• std: Standard deviation (measure of spread)
• min/max: Minimum and maximum values
• 25%, 50%, 75%: Quartiles showing distribution""")

# 5. Target Variable Analysis
doc.add_heading('5. Target Variable Analysis (FLAG)', level=1)
doc.add_paragraph("""The FLAG column is our target variable, indicating whether a transaction is fraudulent (1) or legitimate (0). Understanding the class distribution is crucial for identifying class imbalance.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df['FLAG'].value_counts()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

doc.add_paragraph("""Output:
FLAG
0    7662  (Legitimate transactions - 77.9%)
1    2179  (Fraudulent transactions - 22.1%)

Analysis: The dataset shows a class imbalance where legitimate transactions are significantly more than fraudulent ones. This is an important consideration for model training and evaluation.""")

# 6. Class Distribution Visualization
doc.add_heading('6. Class Distribution Visualization', level=1)
doc.add_paragraph("""A count plot visualizes the distribution of the target variable (FLAG), making it easy to see the proportion of fraudulent vs legitimate transactions.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""sns.countplot(x='FLAG', data=df)
plt.title('Class Distribution (FLAG)')
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# Add image
try:
    doc.add_picture('class_distribution.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("""Screenshot Description:
• X-axis: FLAG values (0 for legitimate, 1 for fraud)
• Y-axis: Count of transactions in each class
• Bar Height Analysis: The legitimate transactions (FLAG=0) bar is significantly taller (~7662 transactions) compared to fraudulent transactions (FLAG=1) bar (~2179 transactions)
• Interpretation: This visual representation clearly shows the class imbalance in the dataset, which is crucial for understanding the modeling challenges and selecting appropriate evaluation metrics""")

# 7. Feature Distribution Analysis
doc.add_heading('7. Feature Distribution Analysis', level=1)
doc.add_paragraph("""Histograms for all numerical features reveal the distribution patterns of each variable. This helps identify skewness, multimodal distributions, and potential data quality issues.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.hist(figsize=(15,12))
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# Add image
try:
    doc.add_picture('feature_distributions.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("""Screenshot Description:
• Multiple subplots showing histogram for each numerical feature (48 features total)
• Distribution Patterns Observed:
  - Right-skewed distributions: Most features show right skewness with most values concentrated on the left
  - Zero-inflated features: Some features have high frequency at zero (e.g., contract creation features)
  - Sparse distributions: Many features have data concentrated in small ranges
  - Multi-modal patterns: Some features show multiple peaks indicating different transaction types
• Quality Issues: Some features show potential data quality issues with extreme outliers
• Fraud Indicator: Features with different distributions for fraud vs legitimate transactions can serve as good predictors""")

# 8. Outlier Detection
doc.add_heading('8. Outlier Detection using Boxplots', level=1)
doc.add_paragraph("""Boxplots are excellent for identifying outliers. The box represents the interquartile range (IQR), with the line inside showing the median, and points beyond the whiskers representing potential outliers.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""plt.figure(figsize=(12,6))
sns.boxplot(data=df.select_dtypes(include=np.number))
plt.xticks(rotation=90)
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# Add image
try:
    doc.add_picture('boxplot_outliers.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("""Screenshot Description:
• Box Representation: Each feature has its own boxplot showing the distribution
• Components Visible:
  - Boxes: Represent the interquartile range (25th to 75th percentile) where 50% of data lies
  - Line in box: Shows the median value
  - Whiskers: Extend to 1.5 × IQR from quartiles
  - Circles/Points: Outliers beyond the whiskers
• Observations:
  - Many features show significant outliers, particularly in Ethereum value columns
  - Some features have extreme outliers reaching up to 10^12
  - Outlier prevalence may indicate fraudulent transactions with unusual patterns
  - Features like total Ether balance show dramatic outliers suggesting whale accounts or contract anomalies
• Significance: These outliers are crucial for fraud detection as fraudulent accounts often exhibit unusual transaction patterns""")

# 9. Correlation Analysis
doc.add_heading('9. Correlation Analysis', level=1)
doc.add_paragraph("""A correlation heatmap shows relationships between numerical features. Strong correlations (positive or negative) indicate that two features move together, which could be important for feature selection.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""numeric_df = df.select_dtypes(include=['int64', 'float64'])
plt.figure(figsize=(14,10))
sns.heatmap(numeric_df.corr(), cmap='coolwarm', linewidths=0.5)
plt.title("Correlation Heatmap (Numeric Features Only)")
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# Add image
try:
    doc.add_picture('correlation_heatmap.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("""Screenshot Description:
• Color Scale: Ranges from deep blue (strong negative correlation, -1) through white (no correlation, 0) to deep red (strong positive correlation, +1)
• Key Observations:
  - Highly Correlated Pairs (Red areas): Features related to the same metric show strong positive correlations
    * Total sent values highly correlated with each other
    * Received transaction metrics highly correlated
    * ERC20 token metrics show strong internal correlations
  - Multicollinearity: Several features show high correlation (>0.9) indicating possible redundancy
  - Weak Correlations with FLAG: Most individual features show weak correlation with FLAG, suggesting fraud is a complex pattern
  - Negative Correlations: Few features show strong negative correlations, indicating most relationships are either positive or neutral
• Implications for Modeling:
  - Feature selection needed to reduce multicollinearity
  - Multiple weak correlators may combine for better prediction
  - Non-linear relationships might be important for fraud detection""")

# 10. Relationship Analysis
doc.add_heading('10. Relationship Analysis', level=1)
doc.add_paragraph("""Scatter plots reveal relationships between pairs of features, with color coding for the target variable (fraud status) to identify patterns specific to fraudulent transactions.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""sns.scatterplot(x='total transactions (including tnx to create contract',
               y='total ether received',
               hue='FLAG',
               data=df,
               alpha=0.6)
plt.title("Total Transactions vs Total Ether Received")
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_run.font.color.rgb = RGBColor(0, 0, 139)

# Add image
try:
    doc.add_picture('relationship_scatter.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("""Screenshot Description:
• Axes:
  - X-axis: Total transactions (including contract creation) - ranges from 0 to ~20,000
  - Y-axis: Total Ether received - ranges from 0 to ~3×10^7
• Color Coding:
  - Blue dots (0): Legitimate transactions - densely clustered in lower ranges
  - Orange dots (1): Fraudulent transactions - show more spread and higher values
• Pattern Analysis:
  - Clustering: Most legitimate transactions cluster in the lower left quadrant
  - Fraud Pattern: Fraudulent transactions show a tendency towards higher transaction counts and Ether values
  - Outliers: Some legitimate transactions show extremely high Ether received (potential whale accounts)
  - Linear Relationship: Weak linear relationship between variables suggesting complex interaction effects
  - Fraud Indicators: Fraudulent transactions don't follow the typical clustering pattern, suggesting distinct behavioral characteristics
• Significance: This visualization shows that transaction behavior is distinctly different between fraudulent and legitimate accounts""")

# 11. Key Findings
doc.add_heading('11. Key Findings and Insights', level=1)

findings_data = [
    ("Dataset Characteristics", "Dataset contains 9,841 transactions with 48 features; a comprehensive collection of on-chain blockchain metrics and transaction patterns."),
    ("Class Distribution", "77.9% legitimate transactions (7,662), 22.1% fraudulent (2,179) - shows moderate class imbalance requiring careful model evaluation metrics."),
    ("Feature Distributions", "Most features are right-skewed with zero-inflation; some features reach extreme values suggesting outliers related to large transactions."),
    ("Outlier Patterns", "Significant outliers exist in Ether value columns; outlier prevalence may correlate with fraudulent activities."),
    ("Multicollinearity Issues", "High correlations between related features (>0.9); feature engineering and selection strategies needed."),
    ("Weak Direct Correlations", "Individual features show weak correlation with FLAG; fraud appears to be a complex pattern requiring multiple features."),
    ("Transaction Behavior Differences", "Fraudulent transactions show distinct patterns: higher transaction counts, different value distributions, and unusual clustering patterns."),
    ("Feature Importance", "Features like total transactions, Ether values, and contract interactions show promise for predictive modeling."),
]

for finding, detail in findings_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(finding + ": ").bold = True
    p.add_run(detail)

doc.add_page_break()

# 12. Recommendations for Next Steps
doc.add_heading('12. Recommendations for Next Steps', level=1)

recommendations = [
    "Handle Class Imbalance: Consider techniques like SMOTE, class weights, or stratified sampling for model training.",
    "Feature Engineering: Create new features combining existing ones; extract behavioral patterns from temporal data.",
    "Feature Selection: Remove highly correlated features to reduce multicollinearity; use statistical tests for feature relevance.",
    "Outlier Handling: Investigate extreme values; consider whether to keep, remove, or transform outliers.",
    "Data Normalization: Apply standardization or normalization to features with different scales for algorithms sensitive to feature magnitude.",
    "Model Selection: Consider ensemble methods, non-linear models, and algorithms robust to imbalanced data.",
    "Evaluation Metrics: Use precision, recall, F1-score, and ROC-AUC instead of accuracy due to class imbalance.",
    "Cross-Validation: Employ stratified k-fold cross-validation to ensure representative train/test splits."
]

for i, rec in enumerate(recommendations, 1):
    doc.add_paragraph(f"{i}. {rec}", style='List Number')

# Conclusion
doc.add_page_break()
doc.add_heading('Conclusion', level=1)
doc.add_paragraph("""This exploratory data analysis has provided comprehensive insights into the Ethereum fraud detection dataset. The visualizations and statistical summaries have revealed:

• The structure and characteristics of blockchain transaction data containing diverse metrics across 48 features
• Clear class imbalance with approximately 3:1 ratio of legitimate to fraudulent transactions
• Complex distribution patterns with right-skewness and zero-inflation in most features
• Significant outliers particularly in Ether value columns that may indicate fraud indicators
• Multicollinearity issues requiring feature selection strategies
• Distinct behavioral patterns separating fraudulent from legitimate transactions

These findings form the foundation for developing effective fraud detection models in subsequent assignments. The insights gained will guide:
• Feature engineering and selection strategies
• Appropriate model selection considering the imbalanced classification problem
• Evaluation metrics selection prioritizing fraud detection capability
• Handling of outliers and extreme values

The EDA has established that while individual features show weak correlation with fraud, the combination of multiple transaction characteristics provides sufficient information for building a robust fraud detection system. The next phase will focus on developing and training machine learning models using these insights.""")

# Save the document
doc.save('Assignment_02_EDA_Report_Enhanced.docx')
print("Enhanced Word document 'Assignment_02_EDA_Report_Enhanced.docx' created successfully!")
