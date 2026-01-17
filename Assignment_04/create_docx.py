from docx import Document

doc = Document()

# Title
doc.add_heading('Assignment 04: Model Deployment, Feedback Collection, and Iterative Improvement', 0)

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_entries = [
    'I. Introduction',
    'II. Deployment Tool Comparison',
    'III. Model Utilization from Assignment 03',
    'IV. Deployment Process',
    'V. Feedback Collection & Analysis',
    'VI. Complete Model Utilization Pipeline',
    'VII. Improvement Plan and Versioning',
    'VIII. Assumptions & Limitations',
    'IX. Conclusion',
    'X. Appendix'
]
for entry in toc_entries:
    doc.add_paragraph(entry)

doc.add_page_break()

doc.add_paragraph('Course: Data Science')
doc.add_paragraph('Class: BSCS-F22')
doc.add_paragraph('Instructor: Mr. Ghulam Ali')
doc.add_paragraph('Student Name: Ahmad Faraz')
doc.add_paragraph('Registration No: 215154')
doc.add_paragraph('Dataset Used: Cleaned_Ethereum_Fraud_Detection.csv')

doc.add_heading('I. Introduction', 1)
doc.add_paragraph('This assignment focuses on transitioning the Ethereum Fraud Detection system from model development to real-world deployment. The trained machine learning model from Assignment 03 is deployed using a lightweight web framework, feedback is collected from users, and an improvement roadmap is proposed based on real usage insights.')

doc.add_heading('II. Deployment Tool Comparison', 1)
doc.add_paragraph('1. Streamlit (Selected Tool)', style='List Number')
doc.add_paragraph('Rapid development and minimal boilerplate', style='List Bullet')
doc.add_paragraph('Ideal for data science demos and ML models', style='List Bullet')
doc.add_paragraph('Built-in UI components', style='List Bullet')
doc.add_paragraph('2. Flask', style='List Number')
doc.add_paragraph('Flexible backend framework', style='List Bullet')
doc.add_paragraph('Requires manual UI and routing', style='List Bullet')
doc.add_paragraph('More setup overhead', style='List Bullet')
doc.add_paragraph('3. FastAPI', style='List Number')
doc.add_paragraph('High performance and async support', style='List Bullet')
doc.add_paragraph('Better for production APIs', style='List Bullet')
doc.add_paragraph('Overkill for academic ML deployment', style='List Bullet')
doc.add_paragraph('Justification: Streamlit was selected due to its simplicity, fast prototyping, and suitability for local and academic deployments.')

doc.add_heading('III. Model Utilization from Assignment 03', 1)
doc.add_paragraph('A. Models Developed in Assignment 03')
doc.add_paragraph('Assignment 03 trained and serialized two main models for fraud detection:')
doc.add_paragraph('1. Logistic Regression Model (logistic_regression_model.pkl)', style='List Number')
doc.add_paragraph('Baseline interpretable model', style='List Bullet')
doc.add_paragraph('Uses scaled features', style='List Bullet')
doc.add_paragraph('Fast predictions with probability outputs', style='List Bullet')
doc.add_paragraph('Suitable for simple, linear fraud patterns', style='List Bullet')
doc.add_paragraph('2. Random Forest Model (random_forest_model.pkl)', style='List Number')
doc.add_paragraph('Ensemble-based model', style='List Bullet')
doc.add_paragraph('Captures non-linear relationships', style='List Bullet')
doc.add_paragraph('Superior performance in Assignment 03', style='List Bullet')
doc.add_paragraph('Better at handling complex fraud patterns', style='List Bullet')
doc.add_paragraph('3. Feature Scaler (scaler.pkl)', style='List Number')
doc.add_paragraph('Fitted StandardScaler from Assignment 03 training data', style='List Bullet')
doc.add_paragraph('Required for Logistic Regression input preprocessing', style='List Bullet')
doc.add_paragraph('Ensures feature consistency across deployment', style='List Bullet')
doc.add_paragraph('B. Assignment 03 Training Process')
doc.add_paragraph('Training Data: 75% of cleaned Ethereum fraud dataset', style='List Bullet')
doc.add_paragraph('Testing Data: 25% with stratified sampling (preserved class imbalance)', style='List Bullet')
doc.add_paragraph('Features Used: Numeric transactional and ERC20-based metrics', style='List Bullet')
doc.add_paragraph('Target Variable: FLAG (1 = Fraudulent, 0 = Legitimate)', style='List Bullet')
doc.add_paragraph('Performance: Random Forest achieved superior ROC-AUC and recall', style='List Bullet')
doc.add_paragraph('C. Model Selection for Deployment')
doc.add_paragraph('Random Forest selected for Production Deployment because:')
doc.add_paragraph('Higher accuracy and ROC-AUC score', style='List Bullet')
doc.add_paragraph('Better recall for detecting fraudulent addresses (critical for security)', style='List Bullet')
doc.add_paragraph('Non-linear pattern recognition capabilities', style='List Bullet')
doc.add_paragraph('Probability outputs for risk scoring', style='List Bullet')

doc.add_heading('IV. Deployment Process', 1)
doc.add_paragraph('The model is deployed locally using Streamlit. Users can input transaction features and receive a fraud risk prediction.')
doc.add_paragraph('Steps to Run Locally:')
doc.add_paragraph('1. Install Streamlit: pip install streamlit', style='List Number')
doc.add_paragraph('2. Save the app code as app.py', style='List Number')
doc.add_paragraph('3. Run: streamlit run app.py', style='List Number')

doc.add_heading('V. Feedback Collection & Analysis', 1)
doc.add_paragraph('The deployed application was shared with at least 15 users. Feedback was collected using a Google Form and stored in a CSV file named Suggestions_Dataset.csv.')
doc.add_paragraph('Feedback fields included:')
doc.add_paragraph('Usability', style='List Bullet')
doc.add_paragraph('Prediction relevance', style='List Bullet')
doc.add_paragraph('Suggestions for improvement', style='List Bullet')

doc.add_heading('VI. Complete Model Utilization Pipeline', 1)
doc.add_paragraph('Includes data flow diagram and model comparison table.')
doc.add_paragraph('Conclusion: Random Forest selected for primary deployment due to superior fraud detection capability (higher recall) and better overall performance.')

doc.add_heading('VII. Improvement Plan and Versioning', 1)
doc.add_paragraph('Version 2.0 Planned Improvements:')
doc.add_paragraph('Add more input features for better accuracy', style='List Bullet')
doc.add_paragraph('Improve UI clarity and tooltips', style='List Bullet')
doc.add_paragraph('Add probability-based risk scoring', style='List Bullet')
doc.add_paragraph('Prioritization: Accuracy and usability improvements were prioritized based on recurring user feedback.')

doc.add_heading('VIII. Assumptions & Limitations', 1)
doc.add_paragraph('Assumptions:')
doc.add_paragraph('Cleaned data represents real-world behavior', style='List Bullet')
doc.add_paragraph('Users input realistic transaction values', style='List Bullet')
doc.add_paragraph('Limitations:')
doc.add_paragraph('Class imbalance may bias predictions', style='List Bullet')
doc.add_paragraph('Model generalization limited to Ethereum network', style='List Bullet')
doc.add_paragraph('Local deployment scalability constraints', style='List Bullet')

doc.add_heading('IX. Conclusion', 1)
doc.add_paragraph('This assignment demonstrated the complete lifecycle of a data science project, from model deployment to feedback-driven improvement. Deployment revealed practical challenges and user expectations, guiding the roadmap for future versions.')

doc.add_heading('X. Appendix', 1)
doc.add_paragraph('Streamlit app code', style='List Bullet')
doc.add_paragraph('Feedback form link', style='List Bullet')
doc.add_paragraph('Screenshots of the deployed interface (attached in report PDF)', style='List Bullet')

doc.save('Assignment_04_Deployment_Report_updated.docx')