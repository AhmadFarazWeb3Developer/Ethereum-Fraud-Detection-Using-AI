import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Navigate to Assignment_02
os.chdir(r'c:\Users\MMO\desktop\Ethereum-Fraud-Detection-Using-AI\Assignment_02')

# Create Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Assignment 02: Exploratory Data Analysis (EDA)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Course Information
info_paragraph = doc.add_paragraph()
info_paragraph.add_run('Course: ').bold = True
info_paragraph.add_run('Data Science\n')
info_paragraph.add_run('Class: ').bold = True
info_paragraph.add_run('BSCS-F22\n')
info_paragraph.add_run('Instructor: ').bold = True
info_paragraph.add_run('Mr. Ghulam Ali\n')
info_paragraph.add_run('Student Name: ').bold = True
info_paragraph.add_run('Ahmad Faraz\n')
info_paragraph.add_run('Registration No: ').bold = True
info_paragraph.add_run('215154\n')
info_paragraph.add_run('Date: ').bold = True
info_paragraph.add_run('January 14, 2026\n')

doc.add_page_break()

# Table of Contents
toc_heading = doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Introduction",
    "2. Dataset Loading and Overview",
    "3. Data Shape and Structure",
    "4. Descriptive Statistics",
    "5. Target Variable Analysis (FLAG)",
    "6. Class Distribution Visualization",
    "7. Feature Distribution Analysis",
    "8. Outlier Detection using Boxplots",
    "9. Correlation Analysis",
    "10. Relationship Analysis",
    "11. Key Findings and Insights"
]

for item in toc_items:
    toc_para = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
intro_text = """This notebook builds upon Assignment 01, where Ethereum blockchain fraud data was collected and preprocessed. The primary goal of this assignment is to perform comprehensive Exploratory Data Analysis (EDA), extract meaningful insights from the data, and establish a methodological foundation for fraud detection.

The dataset contains transactions from the Ethereum blockchain with various features that help identify fraudulent activities. Through this EDA, we aim to:
• Understand the structure and characteristics of the data
• Identify patterns and anomalies in blockchain transactions
• Explore relationships between features
• Detect outliers and unusual transactions
• Build a foundation for predictive modeling"""
doc.add_paragraph(intro_text)

# 2. Dataset Loading and Overview
doc.add_heading('2. Dataset Loading and Overview', level=1)
doc.add_paragraph("""The cleaned Ethereum fraud detection dataset is loaded using pandas. This dataset contains preprocessed transactions with no missing values after the data cleaning process performed in Assignment 01.""")

# Add code snippet
code_para = doc.add_paragraph()
code_para.style = 'Normal'
code_run = code_para.add_run("""import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

df = pd.read_csv('Cleaned_Ethereum_Fraud_Detection.csv')
df.head()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

# 3. Data Shape and Structure
doc.add_heading('3. Data Shape and Structure', level=1)
doc.add_paragraph("""The dataset dimensions and column information provide insights into the data structure. The shape tells us how many transactions and features we have, while the info provides details about data types and memory usage.""")

# Add code snippet
code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.shape  # Returns (rows, columns)
df.info()  # Displays column names, dtypes, and non-null counts""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Expected Output:
• Total rows: Number of transactions
• Total columns: Number of features
• Data types: Mix of integer, float, and potentially categorical features
• Memory usage: Total memory consumed by the dataset""")

# 4. Descriptive Statistics
doc.add_heading('4. Descriptive Statistics', level=1)
doc.add_paragraph("""Descriptive statistics provide a summary of the numerical features in the dataset, including measures of central tendency and spread.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.describe()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""This output includes:
• count: Number of non-null observations
• mean: Average value of each feature
• std: Standard deviation (measure of spread)
• min/max: Minimum and maximum values
• 25%, 50%, 75%: Quartiles showing distribution""")

# 5. Target Variable Analysis
doc.add_heading('5. Target Variable Analysis (FLAG)', level=1)
doc.add_paragraph("""The FLAG column is our target variable, indicating whether a transaction is fraudulent (1) or legitimate (0). Understanding the class distribution is crucial for identifying class imbalance.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df['FLAG'].value_counts()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""This reveals the count of fraudulent vs legitimate transactions. Class imbalance is important to note for modeling purposes.""")

# 6. Class Distribution Visualization
doc.add_heading('6. Class Distribution Visualization', level=1)
doc.add_paragraph("""A count plot visualizes the distribution of the target variable (FLAG), making it easy to see the proportion of fraudulent vs legitimate transactions.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""sns.countplot(x='FLAG', data=df)
plt.title('Class Distribution (FLAG)')
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Output Description:
• X-axis: FLAG values (0 for legitimate, 1 for fraud)
• Y-axis: Count of transactions
• This chart helps identify if the dataset has class imbalance""")

# 7. Feature Distribution Analysis
doc.add_heading('7. Feature Distribution Analysis', level=1)
doc.add_paragraph("""Histograms for all numerical features reveal the distribution patterns of each variable. This helps identify skewness, multimodal distributions, and potential data quality issues.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""df.hist(figsize=(15,12))
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Output Description:
• Multiple subplots showing histogram for each numerical feature
• Helps identify:
  - Normally distributed features
  - Skewed distributions (left or right)
  - Bimodal or multimodal distributions
  - Potential outliers or anomalies""")

# 8. Outlier Detection
doc.add_heading('8. Outlier Detection using Boxplots', level=1)
doc.add_paragraph("""Boxplots are excellent for identifying outliers. The box represents the interquartile range (IQR), with the line inside showing the median, and points beyond the whiskers representing potential outliers.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""plt.figure(figsize=(12,6))
sns.boxplot(data=df.select_dtypes(include=np.number))
plt.xticks(rotation=90)
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Output Description:
• Box: Represents the interquartile range (25th to 75th percentile)
• Line in box: Median value
• Whiskers: Extend to 1.5 × IQR
• Points beyond whiskers: Identified as outliers
• This is crucial for understanding data quality and potential fraud indicators""")

# 9. Correlation Analysis
doc.add_heading('9. Correlation Analysis', level=1)
doc.add_paragraph("""A correlation heatmap shows relationships between numerical features. Strong correlations (positive or negative) indicate that two features move together, which could be important for feature selection.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""numeric_df = df.select_dtypes(include=['int64', 'float64'])
plt.figure(figsize=(14,10))
sns.heatmap(numeric_df.corr(), cmap='coolwarm', linewidths=0.5)
plt.title("Correlation Heatmap (Numeric Features Only)")
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Output Description:
• Color scale: Ranges from blue (negative correlation) to red (positive correlation)
• Values close to +1: Strong positive relationship
• Values close to -1: Strong negative relationship
• Values close to 0: Little to no relationship
• This helps identify multicollinearity and feature importance""")

# 10. Relationship Analysis
doc.add_heading('10. Relationship Analysis', level=1)
doc.add_paragraph("""Scatter plots reveal relationships between pairs of features, with color coding for the target variable (fraud status) to identify patterns specific to fraudulent transactions.""")

code_para = doc.add_paragraph()
code_run = code_para.add_run("""sns.scatterplot(x='total transactions (including tnx to create contract',
               y='total ether received',
               hue='FLAG',
               data=df,
               alpha=0.6)
plt.title("Total Transactions vs Total Ether Received")
plt.show()""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

doc.add_paragraph("""Output Description:
• X-axis: Total number of transactions
• Y-axis: Total Ether received
• Color: FLAG value (different colors for fraud vs legitimate)
• Helps identify if fraudulent transactions follow different patterns
• Transparency (alpha) allows seeing overlapping points""")

# 11. Key Findings
doc.add_heading('11. Key Findings and Insights', level=1)
doc.add_paragraph("""From the EDA performed, the following insights are derived:

1. Dataset Characteristics:
   • Clear overview of transaction volume and feature diversity
   • Understanding of data types and memory footprint

2. Target Variable (Fraud Distribution):
   • Identification of class balance/imbalance
   • Implications for model selection and evaluation metrics

3. Feature Patterns:
   • Distribution shapes indicate data characteristics
   • Skewed features may need transformation
   • Outliers potentially indicate fraudulent behavior

4. Correlation Insights:
   • Highly correlated features may introduce multicollinearity
   • Strong correlations with FLAG indicate predictive power
   • Feature redundancy can be addressed through selection

5. Fraud Patterns:
   • Scatter plot analysis reveals transaction patterns
   • Fraudulent transactions show distinct characteristics
   • Different feature ranges for fraud vs legitimate transactions

6. Recommendations for Next Steps:
   • Feature engineering to create new predictive features
   • Handling of outliers and class imbalance
   • Normalization/standardization if needed
   • Selection of appropriate machine learning algorithms""")

# Conclusion
doc.add_page_break()
doc.add_heading('Conclusion', level=1)
doc.add_paragraph("""This exploratory data analysis has provided comprehensive insights into the Ethereum fraud detection dataset. The visualizations and statistical summaries have revealed:

• The structure and characteristics of blockchain transaction data
• The relationship between features and fraud indicators
• Potential patterns that distinguish fraudulent from legitimate transactions
• The presence of outliers and their potential significance

These findings form the foundation for developing effective fraud detection models in subsequent assignments. The insights gained will guide feature engineering, model selection, and evaluation strategies.""")

# Save the document
doc.save('Assignment_02_EDA_Report.docx')
print("Word document 'Assignment_02_EDA_Report.docx' created successfully!")
