from docx import Document

doc = Document()

# Title
doc.add_heading('App Folder Report: Ethereum Fraud Detection Streamlit Application', 0)

doc.add_paragraph('This report describes the Streamlit application in the App folder, which deploys the machine learning models for Ethereum fraud detection.')

doc.add_heading('Overview', 1)
doc.add_paragraph('The app.py file contains a Streamlit web application that allows users to input Ethereum address features and receive fraud risk predictions using pre-trained models.')

doc.add_heading('Key Components', 1)

doc.add_heading('1. Model Loading', 2)
doc.add_paragraph('The app loads three pickled files:')
doc.add_paragraph('- random_forest_model.pkl: Primary Random Forest model')
doc.add_paragraph('- logistic_regression_model.pkl: Baseline Logistic Regression model')
doc.add_paragraph('- scaler.pkl: StandardScaler for feature normalization')

doc.add_heading('2. Dataset Loading', 2)
doc.add_paragraph('Loads the Cleaned_Ethereum_Fraud_Detection.csv to get feature columns and statistics for input validation.')

doc.add_heading('3. User Interface', 2)
doc.add_paragraph('Features a dark theme with green accents.')
doc.add_paragraph('- Header with title and dataset info')
doc.add_paragraph('- Sidebar with model overview and metrics')
doc.add_paragraph('- Input fields for essential features like transactions sent/received, ETH amounts, ERC20 transactions')

doc.add_heading('4. Prediction Logic', 2)
doc.add_paragraph('Upon clicking "Predict Fraud Risk":')
doc.add_paragraph('- Builds feature vector from user inputs')
doc.add_paragraph('- Random Forest makes primary prediction with probability')
doc.add_paragraph('- Displays result in a colored card (red for high risk, green for low)')
doc.add_paragraph('- Optional comparison with Logistic Regression model in a table')

doc.add_heading('How to Run', 1)
doc.add_paragraph('1. Ensure all required files are in the same directory.')
doc.add_paragraph('2. Install Streamlit: pip install streamlit')
doc.add_paragraph('3. Run: streamlit run app.py')
doc.add_paragraph('4. Access the app in your browser at the provided local URL.')

doc.add_heading('Screenshots', 1)
doc.add_paragraph('Screenshots of the app interface would include:')
doc.add_paragraph('- Main page with input fields')
doc.add_paragraph('- Prediction results display')
doc.add_paragraph('- Sidebar model information')
doc.add_paragraph('(Note: Actual screenshots not included in this DOCX; refer to the running app for visuals.)')

doc.save('App_Report.docx')

print("DOCX created successfully.")